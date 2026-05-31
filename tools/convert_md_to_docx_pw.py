"""
Convert a single Markdown file to DOCX format.

This script takes a markdown file path as a parameter and converts it
into a formatted DOCX document.

Usage:
    python convert_md_to_docx.py <input_markdown_file> [output_docx_file]
    
Example:
    python convert_md_to_docx.py myfile.md
    python convert_md_to_docx.py myfile.md output.docx
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import tempfile
import asyncio
from playwright.async_api import async_playwright


async def render_mermaid_to_image(mermaid_code, output_path):
    """
    Render Mermaid diagram to an image using Playwright.
    
    Args:
        mermaid_code: The Mermaid diagram code
        output_path: Path to save the rendered image
    """
    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <script type="module">
            import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
            mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
        </script>
    </head>
    <body>
        <div class="mermaid">
{mermaid_code}
        </div>
    </body>
    </html>
    """
    
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        await page.set_content(html_template)
        
        # Wait for Mermaid to render
        await page.wait_for_timeout(2000)
        
        # Find the SVG element and take screenshot
        element = await page.query_selector('.mermaid')
        if element:
            await element.screenshot(path=output_path)
        
        await browser.close()

def add_image_to_doc(doc, image_path, alt_text, md_dir):
    """
    Add an image from a file path to the DOCX document.
    
    Args:
        doc: Document object
        image_path: Path to the image (relative or absolute)
        alt_text: Alt text for the image
        md_dir: Directory of the source markdown file (for resolving relative paths)
        
    Returns:
        True if image was added successfully, False otherwise
    """
    # Resolve relative paths against the markdown file's directory
    img_path = Path(image_path)
    if not img_path.is_absolute():
        img_path = md_dir / img_path
    img_path = img_path.resolve()
    
    if not img_path.exists():
        print(f"    Warning: Image not found: {img_path}")
        doc.add_paragraph(f"[Image not found: {alt_text or image_path}]")
        return False
    
    try:
        try:
            from PIL import Image
            img = Image.open(str(img_path))
            img_width, img_height = img.size
            aspect_ratio = img_height / img_width
            
            max_width = Inches(6.5)
            max_height = Inches(8)
            
            width = max_width
            height = width * aspect_ratio
            if height > max_height:
                height = max_height
                width = height / aspect_ratio
            
            doc.add_picture(str(img_path), width=width, height=height)
            print(f"    Added image: {img_path.name} ({img_width}x{img_height}px)")
        except ImportError:
            doc.add_picture(str(img_path), width=Inches(5.5))
            print(f"    Added image: {img_path.name} (PIL not available for sizing)")
        
        # Add caption if alt text is provided
        if alt_text:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(alt_text)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.italic = True
        
        return True
    except Exception as e:
        print(f"    Warning: Failed to add image {image_path}: {e}")
        doc.add_paragraph(f"[Image failed to load: {alt_text or image_path}]")
        return False

def add_bordered_box(doc, text):
    """
    Add a bordered text box/frame for code blocks.
    
    Args:
        doc: Document object
        text: Text content to add in the box
    """
    paragraph = doc.add_paragraph()
    
    # Set paragraph formatting
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    paragraph_format.right_indent = Inches(0.25)
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    
    # Add the text with monospace font
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # Add border to paragraph
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Border size
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '808080')  # Gray color
        pBdr.append(border)
    
    pPr.append(pBdr)
    
    # Add shading (light gray background)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)


def parse_markdown_to_docx(md_file_path, doc, temp_dir):
    """
    Parse a markdown file and add its content to the DOCX document.
    
    Args:
        md_file_path: Path to the markdown file
        doc: Document object to add content to
        temp_dir: Temporary directory for storing generated images
    """
    md_dir = Path(md_file_path).parent.resolve()

    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Check for markdown table
        if line.startswith('|') and '|' in line[1:]:
            # Start or continue table
            if not in_table:
                in_table = True
                table_data = []
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last elements
            table_data.append(cells)
            
            i += 1
            continue
        elif in_table:
            # End of table - render it
            if len(table_data) > 0:
                # Filter out separator rows (rows with only dashes and pipes)
                filtered_rows = [row for row in table_data if not all(set(cell.replace('-', '').replace(':', '').strip()) == set() for cell in row)]
                
                if len(filtered_rows) >= 1:
                    # Create table in document
                    table = doc.add_table(rows=len(filtered_rows), cols=len(filtered_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    # Populate table
                    for row_idx, row_data in enumerate(filtered_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            # Handle bold markdown in cells
                            if '**' in cell_text:
                                paragraph = cell.paragraphs[0]
                                parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                                for part in parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        run = paragraph.add_run(part[2:-2])
                                        run.bold = True
                                    elif part:
                                        paragraph.add_run(part)
                            else:
                                cell.text = cell_text
                            
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    print(f"    Added table with {len(filtered_rows)} rows")
            
            in_table = False
            table_data = []
            # Don't increment i, process this line normally
        
        if re.match(r'^\s*!\[.*?\]\(.*?\)\s*$', line):
            img_match = re.match(r'^\s*!\[(.*?)\]\((.*?)\)\s*$', line)
            if img_match:
                alt_text, image_path = img_match.groups()
                add_image_to_doc(doc, image_path, alt_text, md_dir)
            i += 1
            continue

        # Mermaid code blocks
        if line.startswith('```mermaid'):
            # Collect the mermaid code
            mermaid_code = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                mermaid_code.append(lines[i])
                i += 1
            
            # Render mermaid to image
            mermaid_text = '\n'.join(mermaid_code)
            image_path = temp_dir / f'mermaid_{hash(mermaid_text)}.png'
            
            try:
                # Run async function synchronously
                asyncio.run(render_mermaid_to_image(mermaid_text, str(image_path)))
                
                # Add image to document
                doc.add_picture(str(image_path), width=Inches(6))
                print(f"    Added Mermaid diagram as image")
            except Exception as e:
                print(f"    Warning: Failed to render Mermaid diagram: {e}")
                doc.add_paragraph(f"[Mermaid diagram - rendering failed]")
            
            i += 1
            continue
        
        # Main heading (# )
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
        
        # Subheading (## )
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.runs[0].font.color.rgb = RGBColor(47, 79, 79)
        
        # Sub-subheading (### )
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        
        # Sub-sub-subheading (#### )
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
        
        # Bold text (**text**)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Bullet points (- or *)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            match = re.match(r'^\d+\.\s(.+)', line)
            if match:
                doc.add_paragraph(match.group(1), style='List Number')
        
        # Links [text](url)
        elif '[' in line and '](' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\[.*?\]\(.*?\))', line)
            for part in parts:
                link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if link_match:
                    text, url = link_match.groups()
                    run = p.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                else:
                    p.add_run(part)
        
        # Code blocks (```)
        elif line.startswith('```'):
            # Extract language identifier if present
            lang_match = re.match(r'^```(\w+)?', line)
            language = lang_match.group(1) if lang_match and lang_match.group(1) else None
            
            # Collect code block content
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_content = '\n'.join(code_lines)
            
            # Check if it's a special language (json, c#, etc.)
            if language and language.lower() in ['json', 'csharp', 'c#', 'cs', 'javascript', 'js', 'python', 'py', 'sql', 'xml', 'html', 'css', 'typescript', 'ts']:
                # For specific languages, add a label and use bordered box
                if language.lower() in ['c#', 'cs']:
                    display_lang = 'C#'
                elif language.lower() == 'csharp':
                    display_lang = 'C#'
                elif language.lower() in ['javascript', 'js']:
                    display_lang = 'JavaScript'
                elif language.lower() in ['typescript', 'ts']:
                    display_lang = 'TypeScript'
                elif language.lower() in ['python', 'py']:
                    display_lang = 'Python'
                else:
                    display_lang = language.upper()
                
                # Add language label
                label_p = doc.add_paragraph()
                label_run = label_p.add_run(f"[{display_lang}]")
                label_run.font.size = Pt(8)
                label_run.font.color.rgb = RGBColor(128, 128, 128)
                label_run.italic = True
                label_p.paragraph_format.space_after = Pt(0)
                
                # Add code with syntax highlighting style
                code_p = doc.add_paragraph()
                code_p.paragraph_format.left_indent = Inches(0.25)
                code_p.paragraph_format.space_before = Pt(0)
                code_run = code_p.add_run(code_content)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0, 0, 128)
                print(f"    Added {display_lang} code block")
            else:
                # For generic code blocks (no language or unrecognized language), use bordered box
                add_bordered_box(doc, code_content)
                print(f"    Added generic code block in bordered box")
            
            i += 1
            continue
        
        # Empty lines
        elif not line.strip():
            doc.add_paragraph()
        
        # Regular paragraphs
        else:
            doc.add_paragraph(line)
        
        i += 1


def convert_md_to_docx(input_file, output_file):
    """
    Convert a markdown file to DOCX format.
    
    Args:
        input_file: Path to the input markdown file
        output_file: Path to the output DOCX file
    """
    input_path = Path(input_file)
    
    # Check if input file exists
    if not input_path.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)
    
    # Check if input is a markdown file
    if input_path.suffix.lower() not in ['.md', '.markdown']:
        print(f"Warning: Input file doesn't have a .md or .markdown extension")
    
    # Create document
    doc = Document()
    
    # Create temporary directory for Mermaid images
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        print(f"Converting: {input_file}")
        print(f"Output: {output_file}")
        print("=" * 60)
        
        # Parse and add the markdown content
        parse_markdown_to_docx(input_path, doc, temp_dir)
        
        # Save the document
        doc.save(output_file)
        print(f"\n✓ Successfully created: {output_file}")
    
    except Exception as e:
        print(f"\n✗ Error during conversion: {e}")
        sys.exit(1)
    
    finally:
        # Clean up temporary directory
        import shutil
        if temp_dir.exists():
            shutil.rmtree(temp_dir)


def main():
    """Main entry point."""
    # Check command line arguments
    if len(sys.argv) < 2:
        print("Usage: python convert_md_to_docx.py <input_markdown_file> [output_docx_file]")
        print("\nExample:")
        print("  python convert_md_to_docx.py myfile.md")
        print("  python convert_md_to_docx.py myfile.md output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # Determine output file
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        # Generate output filename based on input filename
        input_path = Path(input_file)
        output_file = str(input_path.with_suffix('.docx'))
    
    # Convert
    convert_md_to_docx(input_file, output_file)


if __name__ == '__main__':
    main()
