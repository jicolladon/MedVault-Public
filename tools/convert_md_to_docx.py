"""
Convert a single Markdown file to DOCX format.

This script takes a markdown file path as a parameter and converts it
into a formatted DOCX document.

Prerequisites:
    Python: 3.6 or higher
    
    Required Python packages:
        - python-docx
        
    Installation command:
        pip install python-docx
        
    Optional (for Mermaid diagram support):
        - Node.js (v14 or higher)
        - npm (comes with Node.js)
        - Mermaid CLI (@mermaid-js/mermaid-cli)
        
    Installation command for Mermaid support:
        npm install -g @mermaid-js/mermaid-cli
        
    Note: After installing Mermaid CLI, restart your terminal or VS Code
          for the PATH changes to take effect.

Usage:
    python convert_md_to_docx.py <input_markdown_file> [output_docx_file] [--mermaid-renderer {mmdc,playwright}]
    
Example:
    python convert_md_to_docx.py myfile.md
    python convert_md_to_docx.py myfile.md output.docx
    python convert_md_to_docx.py myfile.md output.docx --mermaid-renderer playwright
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import tempfile
import subprocess
import shutil
import asyncio
import argparse


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL for the hyperlink
        text: The display text for the hyperlink
        
    Returns:
        The hyperlink run
    """
    # This gets access to the document.xml.rels file and gets a new relationship id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run object
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Apply blue color and underline to the run
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink


def render_mermaid_to_image(mermaid_code, output_path):
    """
    Render Mermaid diagram to an image using local Mermaid CLI (mmdc).
    
    Requires: npm install -g @mermaid-js/mermaid-cli
    
    Args:
        mermaid_code: The Mermaid diagram code
        output_path: Path to save the rendered image
    """
    try:
        # Try to find mmdc command (with .cmd extension on Windows)
        mmdc_cmd = 'mmdc.cmd' if os.name == 'nt' else 'mmdc'
        mmdc_path = shutil.which(mmdc_cmd)
        
        if not mmdc_path:
            # Try alternate name
            mmdc_cmd = 'mmdc'
            mmdc_path = shutil.which(mmdc_cmd)
        
        if not mmdc_path:
            print("    Warning: mmdc (Mermaid CLI) not found in PATH.")
            print("    Install with: npm install -g @mermaid-js/mermaid-cli")
            print("    After installation, restart your terminal or VS Code.")
            return False
        
        # Create temporary input file for mermaid code
        temp_input = Path(output_path).parent / f'temp_mermaid_{hash(mermaid_code)}.mmd'
        
        # Write mermaid code to temp file
        with open(temp_input, 'w', encoding='utf-8') as f:
            f.write(mermaid_code)
        
        # Run mmdc to convert mermaid to PNG
        result = subprocess.run(
            [mmdc_path, '-i', str(temp_input), '-o', str(output_path), '-b', 'transparent'],
            capture_output=True,
            text=True,
            timeout=30,
            shell=True if os.name == 'nt' else False
        )
        
        # Clean up temp input file
        if temp_input.exists():
            temp_input.unlink()
        
        if result.returncode == 0 and Path(output_path).exists():
            return True
        else:
            if result.stderr:
                print(f"    Error from mmdc: {result.stderr}")
            return False
            
    except subprocess.TimeoutExpired:
        print(f"    Error: mmdc command timed out")
        if temp_input.exists():
            temp_input.unlink()
        return False
    except FileNotFoundError:
        print(f"    Warning: mmdc command not found. Install with: npm install -g @mermaid-js/mermaid-cli")
        return False
    except Exception as e:
        print(f"    Error rendering Mermaid: {e}")
        return False


async def render_mermaid_to_image_playwright(mermaid_code, output_path):
    """
    Render Mermaid diagram to an image using Playwright.

    Args:
        mermaid_code: The Mermaid diagram code
        output_path: Path to save the rendered image
    """
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        print("    Warning: Playwright is not installed.")
        print("    Install with: pip install playwright")
        print("    Then install browsers with: playwright install chromium")
        return False

    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <script type="module">
            import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
            mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
        </script>
    </head>
    <body>
        <div class="mermaid">
{mermaid_code}
        </div>
    </body>
    </html>
    """

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.set_content(html_template)

            # Wait for Mermaid to render
            await page.wait_for_timeout(2000)

            # Find the Mermaid element and take screenshot
            element = await page.query_selector('.mermaid')
            if element:
                await element.screenshot(path=output_path)

            await browser.close()

        return Path(output_path).exists()
    except Exception as e:
        print(f"    Error rendering Mermaid with Playwright: {e}")
        return False


def render_mermaid_with_selected_renderer(mermaid_code, output_path, mermaid_renderer='mmdc'):
    """
    Render Mermaid diagram using the selected renderer.

    Args:
        mermaid_code: The Mermaid diagram code
        output_path: Path to save the rendered image
        mermaid_renderer: 'mmdc' (default) or 'playwright'
    """
    if mermaid_renderer == 'playwright':
        return asyncio.run(render_mermaid_to_image_playwright(mermaid_code, output_path))
    return render_mermaid_to_image(mermaid_code, output_path)


def add_image_to_doc(doc, image_path, alt_text, md_dir):
    """
    Add an image from a file path to the DOCX document.
    
    Args:
        doc: Document object
        image_path: Path to the image (relative or absolute)
        alt_text: Alt text for the image
        md_dir: Directory of the source markdown file (for resolving relative paths)
        
    Returns:
        True if image was added successfully, False otherwise
    """
    # Resolve relative paths against the markdown file's directory
    img_path = Path(image_path)
    if not img_path.is_absolute():
        img_path = md_dir / img_path
    img_path = img_path.resolve()
    
    if not img_path.exists():
        print(f"    Warning: Image not found: {img_path}")
        doc.add_paragraph(f"[Image not found: {alt_text or image_path}]")
        return False
    
    try:
        try:
            from PIL import Image
            img = Image.open(str(img_path))
            img_width, img_height = img.size
            aspect_ratio = img_height / img_width
            
            max_width = Inches(6.5)
            max_height = Inches(8)
            
            width = max_width
            height = width * aspect_ratio
            if height > max_height:
                height = max_height
                width = height / aspect_ratio
            
            doc.add_picture(str(img_path), width=width, height=height)
            print(f"    Added image: {img_path.name} ({img_width}x{img_height}px)")
        except ImportError:
            doc.add_picture(str(img_path), width=Inches(5.5))
            print(f"    Added image: {img_path.name} (PIL not available for sizing)")
        
        # Add caption if alt text is provided
        if alt_text:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(alt_text)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.italic = True
        
        return True
    except Exception as e:
        print(f"    Warning: Failed to add image {image_path}: {e}")
        doc.add_paragraph(f"[Image failed to load: {alt_text or image_path}]")
        return False


def add_bordered_box(doc, text):
    """
    Add a bordered text box/frame for code blocks.
    
    Args:
        doc: Document object
        text: Text content to add in the box
    """
    paragraph = doc.add_paragraph()
    
    # Set paragraph formatting
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    paragraph_format.right_indent = Inches(0.25)
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    
    # Add the text with monospace font
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # Add border to paragraph
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Border size
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '808080')  # Gray color
        pBdr.append(border)
    
    pPr.append(pBdr)
    
    # Add shading (light gray background)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)


def add_inline_markdown(paragraph, text):
    """
    Add inline markdown (links and bold) to an existing paragraph.

    Args:
        paragraph: Target paragraph
        text: Inline markdown text
    """
    parts = re.split(r'(\[.*?\]\(.*?\))', text)
    for part in parts:
        link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
        if link_match:
            link_text, url = link_match.groups()
            add_hyperlink(paragraph, url, link_text)
        elif part:
            bold_parts = re.split(r'(\*\*.*?\*\*)', part)
            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**'):
                    run = paragraph.add_run(bold_part[2:-2])
                    run.bold = True
                elif bold_part:
                    paragraph.add_run(bold_part)


def calculate_list_level(indent_text):
    """
    Calculate markdown list nesting level from indentation.

    Markdown commonly uses 2 or 4 spaces (or tabs) for nesting.
    """
    expanded = indent_text.expandtabs(4)
    spaces = len(expanded)
    if spaces < 2:
        return 0
    return min(spaces // 2, 2)


def get_list_style(is_ordered, level):
    """
    Return the Word list style for ordered/bullet list and nesting level.
    """
    if is_ordered:
        styles = ['List Number', 'List Number 2', 'List Number 3']
    else:
        styles = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
    return styles[min(level, len(styles) - 1)]


def parse_markdown_to_docx(md_file_path, doc, temp_dir, mermaid_renderer='mmdc'):
    """
    Parse a markdown file and add its content to the DOCX document.
    
    Args:
        md_file_path: Path to the markdown file
        doc: Document object to add content to
        temp_dir: Temporary directory for storing generated images
    """
    md_dir = Path(md_file_path).parent.resolve()
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Check for markdown table
        if line.startswith('|') and '|' in line[1:]:
            # Start or continue table
            if not in_table:
                in_table = True
                table_data = []
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last elements
            table_data.append(cells)
            
            i += 1
            continue
        elif in_table:
            # End of table - render it
            if len(table_data) > 0:
                # Filter out separator rows (rows with only dashes and pipes)
                filtered_rows = [row for row in table_data if not all(set(cell.replace('-', '').replace(':', '').strip()) == set() for cell in row)]
                
                if len(filtered_rows) >= 1:
                    # Create table in document using the widest row to avoid index errors
                    max_cols = max((len(row) for row in filtered_rows), default=0)
                    if max_cols == 0:
                        in_table = False
                        table_data = []
                        continue
                    table = doc.add_table(rows=len(filtered_rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Populate table
                    for row_idx, row_data in enumerate(filtered_rows):
                        for col_idx in range(max_cols):
                            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
                            cell = table.rows[row_idx].cells[col_idx]
                            paragraph = cell.paragraphs[0]
                            
                            # Handle links in cells
                            if '[' in cell_text and '](' in cell_text:
                                paragraph.clear()
                                parts = re.split(r'(\[.*?\]\(.*?\))', cell_text)
                                for part in parts:
                                    link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                                    if link_match:
                                        text, url = link_match.groups()
                                        add_hyperlink(paragraph, url, text)
                                    elif part:
                                        # Handle bold within non-link parts
                                        if '**' in part:
                                            bold_parts = re.split(r'(\*\*.*?\*\*)', part)
                                            for bold_part in bold_parts:
                                                if bold_part.startswith('**') and bold_part.endswith('**'):
                                                    run = paragraph.add_run(bold_part[2:-2])
                                                    run.bold = True
                                                elif bold_part:
                                                    paragraph.add_run(bold_part)
                                        else:
                                            paragraph.add_run(part)
                            # Handle bold markdown in cells
                            elif '**' in cell_text:
                                paragraph.clear()
                                parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                                for part in parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        run = paragraph.add_run(part[2:-2])
                                        run.bold = True
                                    elif part:
                                        paragraph.add_run(part)
                            else:
                                cell.text = cell_text
                            
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    print(f"    Added table with {len(filtered_rows)} rows")
            
            in_table = False
            table_data = []
            # Don't increment i, process this line normally
        
        # Standalone image line: ![alt text](image_path)
        if re.match(r'^\s*!\[.*?\]\(.*?\)\s*$', line):
            img_match = re.match(r'^\s*!\[(.*?)\]\((.*?)\)\s*$', line)
            if img_match:
                alt_text, image_path = img_match.groups()
                add_image_to_doc(doc, image_path, alt_text, md_dir)
            i += 1
            continue
        
        # Mermaid code blocks
        if line.startswith('```mermaid'):
            # Collect the mermaid code
            mermaid_code = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                mermaid_code.append(lines[i])
                i += 1
            
            # Render mermaid to image
            mermaid_text = '\n'.join(mermaid_code)
            image_path = temp_dir / f'mermaid_{hash(mermaid_text)}.png'
            
            try:
                # Render using selected Mermaid renderer
                if render_mermaid_with_selected_renderer(mermaid_text, str(image_path), mermaid_renderer):
                    # Add image to document with intelligent sizing
                    try:
                        from PIL import Image
                        # Get image dimensions
                        img = Image.open(str(image_path))
                        img_width, img_height = img.size
                        
                        # Calculate aspect ratio
                        aspect_ratio = img_height / img_width
                        
                        # Maximum width is 6.5 inches (leaving margins)
                        max_width = Inches(6.5)
                        max_height = Inches(8)
                        
                        # Calculate dimensions maintaining aspect ratio
                        if img_width > img_height:
                            # Landscape or square
                            width = max_width
                            height = width * aspect_ratio
                            if height > max_height:
                                height = max_height
                                width = height / aspect_ratio
                        else:
                            # Portrait - scale down tall images
                            height = min(Inches(img_height / 96), max_height)
                            width = height / aspect_ratio
                            if width > max_width:
                                width = max_width
                                height = width * aspect_ratio
                        
                        doc.add_picture(str(image_path), width=width, height=height)
                        print(f"    Added Mermaid diagram as image ({img_width}x{img_height}px)")
                    except ImportError:
                        # Fallback if PIL not available
                        doc.add_picture(str(image_path), width=Inches(5.5))
                        print(f"    Added Mermaid diagram as image (PIL not available)")
                else:
                    doc.add_paragraph(f"[Mermaid diagram - rendering failed]")
            except Exception as e:
                print(f"    Warning: Failed to render Mermaid diagram: {e}")
                doc.add_paragraph(f"[Mermaid diagram - rendering failed]")
            
            i += 1
            continue
        
        # Main heading (# )
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
        
        # Subheading (## )
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.runs[0].font.color.rgb = RGBColor(47, 79, 79)
        
        # Sub-subheading (### )
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        
        # Sub-sub-subheading (#### )
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
        
        # Nested bullet/numbered lists with inline markdown support
        elif re.match(r'^([ \t]*)([-+*]|\d+[.)])\s+(.+)$', line):
            match = re.match(r'^([ \t]*)([-+*]|\d+[.)])\s+(.+)$', line)
            indent, marker, content = match.groups()
            level = calculate_list_level(indent)
            is_ordered = bool(re.match(r'^\d+[.)]$', marker))
            
            style_name = get_list_style(is_ordered, level)

            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                fallback = 'List Number' if is_ordered else 'List Bullet'
                p = doc.add_paragraph(style=fallback)

            # Explicitly set the left indent to visually nest the list if Word ignores the style indent
            if level > 0:
                p.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))

            add_inline_markdown(p, content)

        # Bold text (**text**)
        elif '**' in line:
            p = doc.add_paragraph()
            add_inline_markdown(p, line)
        
        # Lines containing inline images mixed with text
        elif '![' in line and '](' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(!\[.*?\]\(.*?\))', line)
            for part in parts:
                img_match = re.match(r'!\[(.*?)\]\((.*?)\)', part)
                if img_match:
                    alt_text, image_path = img_match.groups()
                    add_image_to_doc(doc, image_path, alt_text, md_dir)
                elif part:
                    p.add_run(part)
        
        # Links in regular text [text](url)
        elif '[' in line and '](' in line:
            p = doc.add_paragraph()
            add_inline_markdown(p, line)
        
        # Code blocks (```)
        elif line.startswith('```'):
            # Extract language identifier if present
            lang_match = re.match(r'^```(\w+)?', line)
            language = lang_match.group(1) if lang_match and lang_match.group(1) else None
            
            # Collect code block content
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_content = '\n'.join(code_lines)
            
            # Check if it's a special language (json, c#, etc.)
            if language and language.lower() in ['json', 'csharp', 'c#', 'cs', 'javascript', 'js', 'python', 'py', 'sql', 'xml', 'html', 'css', 'typescript', 'ts']:
                # For specific languages, add a label and use bordered box
                if language.lower() in ['c#', 'cs']:
                    display_lang = 'C#'
                elif language.lower() == 'csharp':
                    display_lang = 'C#'
                elif language.lower() in ['javascript', 'js']:
                    display_lang = 'JavaScript'
                elif language.lower() in ['typescript', 'ts']:
                    display_lang = 'TypeScript'
                elif language.lower() in ['python', 'py']:
                    display_lang = 'Python'
                else:
                    display_lang = language.upper()
                
                # Add language label
                label_p = doc.add_paragraph()
                label_run = label_p.add_run(f"[{display_lang}]")
                label_run.font.size = Pt(8)
                label_run.font.color.rgb = RGBColor(128, 128, 128)
                label_run.italic = True
                label_p.paragraph_format.space_after = Pt(0)
                
                # Add code with syntax highlighting style
                code_p = doc.add_paragraph()
                code_p.paragraph_format.left_indent = Inches(0.25)
                code_p.paragraph_format.space_before = Pt(0)
                code_run = code_p.add_run(code_content)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0, 0, 128)
                print(f"    Added {display_lang} code block")
            else:
                # For generic code blocks (no language or unrecognized language), use bordered box
                add_bordered_box(doc, code_content)
                print(f"    Added generic code block in bordered box")
            
            i += 1
            continue
        
        # Empty lines
        elif not line.strip():
            doc.add_paragraph()
        
        # Regular paragraphs
        else:
            doc.add_paragraph(line)
        
        i += 1


def convert_md_to_docx(input_file, output_file, mermaid_renderer='mmdc'):
    """
    Convert a markdown file to DOCX format.
    
    Args:
        input_file: Path to the input markdown file
        output_file: Path to the output DOCX file
    """
    input_path = Path(input_file)
    
    # Check if input file exists
    if not input_path.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)
    
    # Check if input is a markdown file
    if input_path.suffix.lower() not in ['.md', '.markdown']:
        print(f"Warning: Input file doesn't have a .md or .markdown extension")
    
    # Create document
    doc = Document()
    
    # Create temporary directory for Mermaid images
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        print(f"Converting: {input_file}")
        print(f"Output: {output_file}")
        print("=" * 60)
        
        # Parse and add the markdown content
        parse_markdown_to_docx(input_path, doc, temp_dir, mermaid_renderer)
        
        # Save the document
        doc.save(output_file)
        print(f"\n✓ Successfully created: {output_file}")
    
    except Exception as e:
        print(f"\n✗ Error during conversion: {e}")
        sys.exit(1)
    
    finally:
        # Clean up temporary directory
        import shutil
        if temp_dir.exists():
            shutil.rmtree(temp_dir)


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='Convert a Markdown file to DOCX format.'
    )
    parser.add_argument('input_markdown_file', help='Path to the input markdown file')
    parser.add_argument('output_docx_file', nargs='?', help='Path to the output DOCX file')
    parser.add_argument(
        '--mermaid-renderer',
        choices=['mmdc', 'playwright'],
        default='mmdc',
        help='Mermaid renderer to use (default: mmdc)'
    )

    args = parser.parse_args()

    input_file = args.input_markdown_file
    output_file = args.output_docx_file

    if not output_file:
        input_path = Path(input_file)
        output_file = str(input_path.with_suffix('.docx'))

    convert_md_to_docx(input_file, output_file, args.mermaid_renderer)


if __name__ == '__main__':
    main()
